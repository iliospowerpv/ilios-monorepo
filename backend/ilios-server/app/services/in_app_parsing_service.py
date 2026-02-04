"""In-App AI Parsing Service

Replaces external cloud function for AI document parsing.
Uses Replit AI Integrations (OpenAI) for text extraction and field parsing.

Flow:
1. Download file bytes from Replit Object Storage
2. Validate file size and page count (guardrails)
3. Extract text from PDF/DOCX using pypdf/python-docx
4. Validate extracted text quality (min chars threshold)
5. Truncate text if exceeds max chars limit
6. Build prompt using ExtractionPipelineService registry
7. Call OpenAI API for extraction
8. Parse and store results in AIParsingResult

Uses gpt-5.2 model via Replit AI Integrations (no API key required, billed to credits).
"""

import json
import logging
import os
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from enum import Enum
from io import BytesIO
from typing import Optional, Tuple

from sqlalchemy.orm import Session
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception

from app.crud.ai_parsing_result import AIParsingResultCRUD
from app.helpers.files.storage_service import get_storage_service
from app.models.file import File as FileModel, FileParsingStatuses
from app.services.extraction_pipeline_service import ExtractionPipelineService
from app.settings import settings

logger = logging.getLogger(__name__)


class ParsingReasonCode(str, Enum):
    """Machine-readable reason codes for parsing failures."""
    FILE_TOO_LARGE = "file_too_large"
    TOO_MANY_PAGES = "too_many_pages"
    INSUFFICIENT_TEXT_EXTRACTED = "insufficient_text_extracted"
    UNSUPPORTED_FILE_TYPE = "unsupported_file_type"
    TEXT_EXTRACTION_FAILED = "text_extraction_failed"
    LLM_CALL_FAILED = "llm_call_failed"
    NO_EXTRACTION_CONFIG = "no_extraction_config"
    STORAGE_ERROR = "storage_error"


@dataclass
class ExtractionResult:
    """Result of text extraction with metadata."""
    text: str
    char_count: int
    word_count: int
    page_count: Optional[int]
    was_truncated: bool
    truncated_char_count: Optional[int] = None


class ParsingGuardrailError(Exception):
    """Exception raised when a parsing guardrail is violated."""
    def __init__(self, reason_code: ParsingReasonCode, message: str, details: Optional[dict] = None):
        self.reason_code = reason_code
        self.message = message
        self.details = details or {}
        super().__init__(f"[{reason_code.value}] {message}")
    
    def formatted_error(self) -> str:
        """Return error message with reason code prefix."""
        return f"[{self.reason_code.value}] {self.message}"


def is_rate_limit_error(exception: BaseException) -> bool:
    """Check if the exception is a rate limit or quota violation error."""
    error_msg = str(exception)
    return (
        "429" in error_msg
        or "RATELIMIT_EXCEEDED" in error_msg
        or "quota" in error_msg.lower()
        or "rate limit" in error_msg.lower()
        or (hasattr(exception, "status_code") and exception.status_code == 429)
    )


def validate_file_size(file_bytes: bytes, max_size_mb: int) -> None:
    """Validate file size against maximum limit."""
    file_size_mb = len(file_bytes) / (1024 * 1024)
    if file_size_mb > max_size_mb:
        raise ParsingGuardrailError(
            reason_code=ParsingReasonCode.FILE_TOO_LARGE,
            message=f"File size ({file_size_mb:.1f} MB) exceeds maximum allowed ({max_size_mb} MB)",
            details={"file_size_mb": round(file_size_mb, 2), "max_size_mb": max_size_mb}
        )


def validate_pdf_page_count(pdf_reader, max_pages: int) -> int:
    """Validate PDF page count and return the count."""
    page_count = len(pdf_reader.pages)
    if page_count > max_pages:
        raise ParsingGuardrailError(
            reason_code=ParsingReasonCode.TOO_MANY_PAGES,
            message=f"PDF has {page_count} pages, exceeds maximum allowed ({max_pages} pages)",
            details={"page_count": page_count, "max_pages": max_pages}
        )
    return page_count


def extract_text_from_pdf(file_bytes: bytes, max_pages: Optional[int] = None) -> Tuple[str, int]:
    """Extract text from PDF bytes using pypdf.
    
    Returns:
        Tuple of (extracted_text, page_count)
    """
    try:
        from pypdf import PdfReader
        pdf_reader = PdfReader(BytesIO(file_bytes))
        
        if max_pages is not None:
            page_count = validate_pdf_page_count(pdf_reader, max_pages)
        else:
            page_count = len(pdf_reader.pages)
        
        text_parts = []
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
        return ("\n\n".join(text_parts), page_count)
    except ParsingGuardrailError:
        raise
    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise ParsingGuardrailError(
            reason_code=ParsingReasonCode.TEXT_EXTRACTION_FAILED,
            message=f"PDF text extraction failed: {e}",
            details={"error": str(e)}
        )


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(BytesIO(file_bytes))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise ParsingGuardrailError(
            reason_code=ParsingReasonCode.TEXT_EXTRACTION_FAILED,
            message=f"DOCX text extraction failed: {e}",
            details={"error": str(e)}
        )


def validate_extracted_text(text: str, min_chars: int) -> None:
    """Validate that extracted text meets minimum quality threshold."""
    if len(text) < min_chars:
        raise ParsingGuardrailError(
            reason_code=ParsingReasonCode.INSUFFICIENT_TEXT_EXTRACTED,
            message=f"Extracted text ({len(text)} chars) below minimum threshold ({min_chars} chars). "
                    "This may be a scanned document requiring OCR.",
            details={"char_count": len(text), "min_chars": min_chars}
        )


def truncate_text_if_needed(text: str, max_chars: int) -> Tuple[str, bool, Optional[int]]:
    """Truncate text to maximum chars if needed.
    
    Returns:
        Tuple of (text, was_truncated, truncated_char_count)
    """
    if len(text) <= max_chars:
        return (text, False, None)
    
    truncated_count = len(text) - max_chars
    truncated_text = text[:max_chars]
    
    newline_pos = truncated_text.rfind("\n", max_chars - 500, max_chars)
    if newline_pos > 0:
        truncated_text = truncated_text[:newline_pos]
    
    logger.warning(
        f"Text truncated from {len(text)} to {len(truncated_text)} chars "
        f"(removed {truncated_count} chars)"
    )
    
    return (truncated_text, True, truncated_count)


def extract_text_with_guardrails(
    file_bytes: bytes,
    filename: str,
    max_file_size_mb: Optional[int] = None,
    max_pdf_pages: Optional[int] = None,
    min_text_chars: Optional[int] = None,
    max_chars_to_llm: Optional[int] = None,
) -> ExtractionResult:
    """Extract text from file with full guardrail validation.
    
    Args:
        file_bytes: Raw file bytes
        filename: Original filename (for extension detection)
        max_file_size_mb: Max file size in MB (None = use settings default)
        max_pdf_pages: Max PDF pages (None = use settings default)
        min_text_chars: Min extracted chars (None = use settings default)
        max_chars_to_llm: Max chars to send to LLM (None = use settings default)
    
    Returns:
        ExtractionResult with text and metadata
    
    Raises:
        ParsingGuardrailError: If any guardrail is violated
    """
    max_file_size_mb = max_file_size_mb or settings.parsing_max_file_size_mb
    max_pdf_pages = max_pdf_pages or settings.parsing_max_pdf_pages
    min_text_chars = min_text_chars or settings.parsing_min_text_chars
    max_chars_to_llm = max_chars_to_llm or settings.parsing_max_chars_to_llm
    
    validate_file_size(file_bytes, max_file_size_mb)
    
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    page_count = None
    
    if ext == "pdf":
        text, page_count = extract_text_from_pdf(file_bytes, max_pdf_pages)
    elif ext in ("docx", "doc"):
        text = extract_text_from_docx(file_bytes)
    else:
        raise ParsingGuardrailError(
            reason_code=ParsingReasonCode.UNSUPPORTED_FILE_TYPE,
            message=f"Unsupported file type: {ext}",
            details={"extension": ext, "supported": ["pdf", "docx", "doc"]}
        )
    
    validate_extracted_text(text, min_text_chars)
    
    text, was_truncated, truncated_char_count = truncate_text_if_needed(text, max_chars_to_llm)
    
    return ExtractionResult(
        text=text,
        char_count=len(text),
        word_count=len(text.split()),
        page_count=page_count,
        was_truncated=was_truncated,
        truncated_char_count=truncated_char_count,
    )


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Extract text from file bytes based on file extension.
    
    DEPRECATED: Use extract_text_with_guardrails for full validation.
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        text, _ = extract_text_from_pdf(file_bytes)
        return text
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type for text extraction: {ext}")


class InAppParsingService:
    """Service for AI document parsing using OpenAI via Replit AI Integrations."""

    def __init__(self, db_session: Session):
        self.db = db_session
        self.pipeline_service = ExtractionPipelineService(db_session)
        self.ai_results_crud = AIParsingResultCRUD(db_session)
        self._openai_client = None

    @property
    def openai_client(self):
        """Lazy-initialize OpenAI client using Replit AI Integrations."""
        if self._openai_client is None:
            try:
                from openai import OpenAI
                api_key = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
                base_url = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")
                if not api_key or not base_url:
                    raise ValueError(
                        "AI_INTEGRATIONS_OPENAI_API_KEY and AI_INTEGRATIONS_OPENAI_BASE_URL must be set. "
                        "Install the OpenAI integration via Replit AI Integrations."
                    )
                self._openai_client = OpenAI(api_key=api_key, base_url=base_url)
                logger.info("OpenAI client initialized via Replit AI Integrations")
            except ImportError as e:
                raise RuntimeError("openai package not installed") from e
        return self._openai_client

    def download_file_bytes(self, file: FileModel) -> bytes:
        """Download file bytes from storage."""
        storage_service = get_storage_service()
        if file.storage_key:
            logger.info(f"Downloading file from Replit storage: {file.storage_key}")
            return storage_service.download_bytes(file.storage_key)
        elif file.filepath:
            logger.warning(f"File {file.id} uses legacy GCS path - not supported in Replit-native mode")
            raise ValueError(
                f"File {file.id} uses legacy GCS storage. "
                "Re-upload the file to use Replit Object Storage."
            )
        else:
            raise ValueError(f"File {file.id} has no storage_key or filepath")

    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=2, max=30),
        retry=retry_if_exception(is_rate_limit_error),
        reraise=True,
    )
    def call_llm(
        self,
        system_prompt: str,
        user_prompt: str,
        model_name: str = "gpt-5.2",
        max_tokens: int = 8192,
    ) -> dict:
        """Call OpenAI API with retry logic for rate limits.
        
        In test mode (LLM_STUB_ENABLED=true), uses deterministic stub instead of real API.
        """
        from app.services.llm_stub import get_llm_stub
        
        stub = get_llm_stub()
        if stub is not None:
            logger.info("Using LLM stub for testing")
            return stub.call(system_prompt, user_prompt)
        
        # the newest OpenAI model is "gpt-5" which was released August 7, 2025.
        # do not change this unless explicitly requested by the user
        model = model_name if model_name else "gpt-5.2"
        if model.startswith("gpt-4") and not model.startswith("gpt-4.1"):
            model = "gpt-5.2"
        logger.info(f"Calling LLM with model: {model}")
        response = self.openai_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            response_format={"type": "json_object"},
            max_completion_tokens=max_tokens,
        )
        content = response.choices[0].message.content
        if not content:
            raise ValueError("Empty response from LLM")
        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse LLM response as JSON: {content[:500]}")
            raise ValueError(f"LLM response is not valid JSON: {e}")

    def parse_file(
        self,
        file: FileModel,
        ai_result_id: int,
        document_type_name: str,
        correlation_id: Optional[str] = None,
    ) -> dict:
        """
        Parse a file and update the AIParsingResult record.

        Args:
            file: The File model instance
            ai_result_id: ID of the AIParsingResult record to update
            document_type_name: Document type name for extraction config lookup
            correlation_id: Optional correlation ID for logging/tracing

        Returns:
            dict with parsed result and metadata
        """
        correlation_id = correlation_id or str(uuid.uuid4())[:8]
        log_prefix = f"[{correlation_id}] File {file.id}"
        logger.info(f"{log_prefix} Starting in-app parsing for document type: {document_type_name}")
        start_time = datetime.now(timezone.utc)
        retries = 0
        extraction_metadata = {}

        try:
            try:
                file_bytes = self.download_file_bytes(file)
            except Exception as e:
                raise ParsingGuardrailError(
                    reason_code=ParsingReasonCode.STORAGE_ERROR,
                    message=f"Failed to download file: {e}",
                    details={"file_id": file.id}
                )
            
            logger.info(f"{log_prefix} Downloaded {len(file_bytes)} bytes")
            
            extraction_result = extract_text_with_guardrails(
                file_bytes=file_bytes,
                filename=file.filename,
            )
            
            extraction_metadata = {
                "char_count": extraction_result.char_count,
                "word_count": extraction_result.word_count,
                "page_count": extraction_result.page_count,
                "was_truncated": extraction_result.was_truncated,
                "truncated_char_count": extraction_result.truncated_char_count,
            }
            
            logger.info(
                f"{log_prefix} Extracted {extraction_result.char_count} chars, "
                f"{extraction_result.word_count} words, "
                f"pages={extraction_result.page_count}, "
                f"truncated={extraction_result.was_truncated}"
            )
            
            prompt_data = self.pipeline_service.build_extraction_prompt(
                document_type_name, extraction_result.text
            )
            if not prompt_data:
                raise ParsingGuardrailError(
                    reason_code=ParsingReasonCode.NO_EXTRACTION_CONFIG,
                    message=f"No extraction config found for document type: {document_type_name}",
                    details={"document_type": document_type_name}
                )
            
            logger.info(
                f"{log_prefix} Built extraction prompt with "
                f"doc_type_id={prompt_data['metadata']['document_type_id']}, "
                f"schema_v={prompt_data['metadata']['schema_version_id']}, "
                f"prompt_v={prompt_data['metadata']['prompt_template_id']}"
            )
            
            try:
                parsed_result = self.call_llm(
                    system_prompt=prompt_data["system_prompt"],
                    user_prompt=prompt_data["user_prompt"],
                    model_name=prompt_data["model_config"]["model_name"],
                    max_tokens=prompt_data["model_config"]["max_tokens"] or 8192,
                )
            except Exception as llm_error:
                raise ParsingGuardrailError(
                    reason_code=ParsingReasonCode.LLM_CALL_FAILED,
                    message=f"LLM extraction failed: {llm_error}",
                    details={"error": str(llm_error)}
                )
            
            logger.info(f"{log_prefix} LLM extraction completed, {len(parsed_result)} fields extracted")
            end_time = datetime.now(timezone.utc)
            
            full_metadata = {**prompt_data["metadata"], **extraction_metadata}
            
            self.ai_results_crud.update_by_id(ai_result_id, {
                "status": FileParsingStatuses.completed,
                "parsed_result": parsed_result,
                "end_time": end_time,
                "document_type_id": prompt_data["metadata"]["document_type_id"],
                "schema_version_id": prompt_data["metadata"]["schema_version_id"],
                "prompt_template_id": prompt_data["metadata"]["prompt_template_id"],
                "retries": retries,
                "error_message": None,
            })
            
            logger.info(f"{log_prefix} Parsing completed successfully in {(end_time - start_time).total_seconds():.2f}s")
            
            return {
                "status": "completed",
                "parsed_result": parsed_result,
                "metadata": full_metadata,
                "duration_seconds": (end_time - start_time).total_seconds(),
            }

        except ParsingGuardrailError as e:
            end_time = datetime.now(timezone.utc)
            error_msg = e.formatted_error()
            logger.error(f"{log_prefix} Guardrail violation: {error_msg}")
            self.ai_results_crud.update_by_id(ai_result_id, {
                "status": FileParsingStatuses.processing_failed,
                "end_time": end_time,
                "error_message": error_msg,
                "retries": retries,
            })
            raise
        except Exception as e:
            end_time = datetime.now(timezone.utc)
            error_msg = str(e)[:500]
            logger.error(f"{log_prefix} Parsing failed: {error_msg}")
            self.ai_results_crud.update_by_id(ai_result_id, {
                "status": FileParsingStatuses.processing_failed,
                "end_time": end_time,
                "error_message": error_msg,
                "retries": retries,
            })
            raise

    def check_openai_available(self) -> bool:
        """Check if OpenAI integration is properly configured.
        
        Validates both environment variables and ability to instantiate client.
        """
        api_key = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
        base_url = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")
        if not api_key or not base_url:
            return False
        try:
            from openai import OpenAI
            OpenAI(api_key=api_key, base_url=base_url)
            return True
        except Exception as e:
            logger.warning(f"OpenAI client instantiation failed: {e}")
            return False


def get_in_app_parsing_service(db_session: Session) -> InAppParsingService:
    """Factory function for dependency injection."""
    return InAppParsingService(db_session)
